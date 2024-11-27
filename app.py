# Version 1.2.2
# Changes:
# - Added CombatStats tracking
# - Improved enemy and damage counting
# - Enhanced fact-based summarization
# - Expanded DnD terminology allowlist

import streamlit as st
import openai
from docx import Document
from io import BytesIO
import re
from typing import Dict, List, Set, Union
from dataclasses import dataclass
from collections import defaultdict

@dataclass
class CombatStats:
    enemies: Dict[str, int] = None
    damage_dealt: Dict[str, List[int]] = None
    damage_taken: Dict[str, List[int]] = None
    kills: Dict[str, int] = None
    
    def __post_init__(self):
        self.enemies = defaultdict(int)
        self.damage_dealt = defaultdict(list)
        self.damage_taken = defaultdict(list)
        self.kills = defaultdict(int)

class DnDSanitizer:
    def __init__(self):
        # Core D&D terms that should never be filtered
        self.dnd_allowlist = {
            'dragon', 'demon', 'devil', 'undead', 'zombie', 'kobold', 'orc',
            'spell', 'magic', 'curse', 'ritual', 'enchant',
            'blood', 'corpse', 'wound', 'kill', 'dead', 'death'
        }
        self.profanity_set = self._load_profanity_list()
        
    def _load_profanity_list(self) -> Set[str]:
        # TODO: Replace with actual profanity list import
        return {'badword1', 'badword2'}
        
    def sanitize(self, text: str) -> str:
        words = text.split()
        sanitized_words = []
        
        for word in words:
            clean_word = word.lower().strip('.,!?')
            if clean_word in self.profanity_set and clean_word not in self.dnd_allowlist:
                sanitized_words.append('[REMOVED]')
            else:
                sanitized_words.append(word)
                
        return ' '.join(sanitized_words)
        
    def is_safe(self, phrase: str) -> bool:
        words = phrase.lower().split()
        return not any(word.strip('.,!?') in self.profanity_set 
                      for word in words 
                      if word not in self.dnd_allowlist)

class SessionAnalyzer:
    def __init__(self, party_members: Dict[str, str]):
        self.party_members = party_members
        self.combat_stats = CombatStats()
        
    def extract_damage(self, text: str) -> None:
        """Extract damage numbers and associate with characters"""
        damage_pattern = r'(\d+)\s*(?:points? of)?\s*damage'
        for line in text.split('\n'):
            damage_matches = re.finditer(damage_pattern, line)
            for match in damage_matches:
                damage = int(match.group(1))
                # Associate damage with character based on context
                if any(member.lower() in line.lower() for member in self.party_members):
                    for member in self.party_members:
                        if member.lower() in line.lower():
                            if 'takes' in line.lower() or 'receives' in line.lower():
                                self.combat_stats.damage_taken[member].append(damage)
                            else:
                                self.combat_stats.damage_dealt[member].append(damage)
                            break
    
    def track_enemies(self, text: str) -> None:
        """Track enemy appearances and deaths"""
        enemy_types = {'kobold', 'orc', 'goblin', 'dragon', 'demon', 'undead'}  # Expand as needed
        for line in text.split('\n'):
            line_lower = line.lower()
            # Count enemies
            for enemy in enemy_types:
                if enemy in line_lower:
                    count_pattern = r'(\d+)\s+' + enemy
                    matches = re.finditer(count_pattern, line_lower)
                    for match in matches:
                        self.combat_stats.enemies[enemy] += int(match.group(1))
            
            # Track kills
            if 'killed' in line_lower or 'defeated' in line_lower:
                for enemy in enemy_types:
                    if enemy in line_lower:
                        kill_pattern = r'(\d+)\s+' + enemy
                        matches = re.finditer(kill_pattern, line_lower)
                        for match in matches:
                            self.combat_stats.kills[enemy] += int(match.group(1))

def chunk_text(text: str, chunk_size: int = 15000) -> List[str]:
    """Split text into chunks while preserving context"""
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = []
    current_length = 0
    
    for para in paragraphs:
        para_length = len(para)
        
        if current_length + para_length > chunk_size and current_chunk:
            chunks.append('\n\n'.join(current_chunk))
            current_chunk = [para]
            current_length = para_length
        else:
            current_chunk.append(para)
            current_length += para_length
    
    if current_chunk:
        chunks.append('\n\n'.join(current_chunk))
    
    return chunks

def generate_summary(text: str, client: openai.OpenAI, analyzer: SessionAnalyzer) -> str:
    """Generate accurate summary based on tracked stats"""
    chunks = chunk_text(text)
    summaries = []
    
    # Process each chunk
    for chunk in chunks:
        analyzer.extract_damage(chunk)
        analyzer.track_enemies(chunk)
        
        # Generate summary with accurate stats
        response = client.chat.completions.create(
            model="gpt-4",
            temperature=0.2,
            messages=[
                {"role": "system", "content": """Generate a D&D session summary with strict accuracy:
- Only include confirmed enemy counts and types
- Use exact damage numbers from combat
- Reference specific character actions and outcomes
- Do not add speculative or decorative details
- Format as structured sections (Combat, Exploration, etc.)
- Include only events explicitly described in the text"""},
                {"role": "user", "content": f"""Summarize this session chunk using these verified statistics:
Enemy Counts: {dict(analyzer.combat_stats.enemies)}
Confirmed Kills: {dict(analyzer.combat_stats.kills)}
Damage Dealt: {dict(analyzer.combat_stats.damage_dealt)}
Damage Taken: {dict(analyzer.combat_stats.damage_taken)}

Session Text:
{chunk}"""}
            ]
        )
        summaries.append(response.choices[0].message.content)
    
    # Combine summaries if needed
    if len(summaries) > 1:
        combined = "\n\n".join(summaries)
        response = client.chat.completions.create(
            model="gpt-4",
            temperature=0.2,
            messages=[
                {"role": "system", "content": "Combine these summaries into one cohesive narrative while maintaining numerical accuracy and factual details."},
                {"role": "user", "content": combined}
            ]
        )
        return response.choices[0].message.content
    
    return summaries[0]

def create_docx(summary: str) -> BytesIO:
    """Create Word document from summary"""
    doc = Document()
    doc.add_heading('D&D Session Summary', 0)
    doc.add_paragraph(summary)
    
    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)
    return docx_file

# Configure Streamlit page
st.set_page_config(page_title="D&D Session Summarizer")
st.title("D&D Session Summarizer")

# Initialize OpenAI client
client = openai.OpenAI(api_key=st.secrets["openai"]["openai_api_key"])

# Party configuration
PARTY_MEMBERS = {
    "Mike": "Rogue",
    "Justin": "Fighter",
    "Dave": "Barbarian",
    "Steve": "Sorcerer",
    "Chris": "Bard",
    "Geoff": "Dungeon Master"
}

# Initialize analyzer
analyzer = SessionAnalyzer(PARTY_MEMBERS)

# File uploader
uploaded_file = st.file_uploader("Upload your D&D session transcript", type=['txt'])

if uploaded_file:
    if st.button('Generate Summary'):
        try:
            text = uploaded_file.read().decode('utf-8')
            
            # Process and summarize
            with st.spinner('Analyzing session...'):
                sanitizer = DnDSanitizer()
                cleaned_text = sanitizer.sanitize(text)
                summary = generate_summary(cleaned_text, client, analyzer)
            
            # Display results
            st.subheader("Summary Statistics:")
            st.write("Enemy Encounters:", dict(analyzer.combat_stats.enemies))
            st.write("Confirmed Kills:", dict(analyzer.combat_stats.kills))
            
            st.subheader("Summary:")
            st.write(summary)
            
            # Create download
            docx_file = create_docx(summary)
            st.download_button(
                label="Download Summary as DOCX",
                data=docx_file,
                file_name=f"summary_{uploaded_file.name.replace('.txt', '.docx')}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        except Exception as e:
            st.error(f"Error: {str(e)}")
